#!/usr/bin/env python3
"""Gera o Word de kickoff + especificação para a equipe."""

from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn, nsmap
from docx.shared import Cm, Inches, Pt, RGBColor

NAVY = RGBColor(0x0A, 0x25, 0x40)
GOLD = RGBColor(0xC8, 0x9D, 0x66)
SLATE = RGBColor(0x33, 0x41, 0x55)
MUTED = RGBColor(0x47, 0x55, 0x69)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
ROW_ALT = "F8F9FA"
HEADER_BG = "0A2540"
GOLD_BG = "C89D66"
GREEN_BG = "ECFDF5"


def set_run(run, *, size=11, bold=False, color=SLATE, font="Calibri"):
    run.font.name = font
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font)
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = color


def shade_cell(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color)
    shd.set(qn("w:val"), "clear")
    tcPr.append(shd)


def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for edge, val in kwargs.items():
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), val.get("val", "single"))
        el.set(qn("w:sz"), val.get("sz", "4"))
        el.set(qn("w:color"), val.get("color", "E2E8F0"))
        tcBorders.append(el)
    tcPr.append(tcBorders)


def cell_text(cell, text, *, bold=False, color=SLATE, size=10, align="left"):
    cell.text = ""
    p = cell.paragraphs[0]
    if align == "center":
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text)
    set_run(run, size=size, bold=bold, color=color)


def add_heading_styled(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = NAVY if level == 1 else RGBColor(0x1E, 0x3A, 0x5F)
        run.font.name = "Calibri"
    return h


def add_p(doc, text, *, size=11, bold=False, color=SLATE, space_after=8):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run(text)
    set_run(run, size=size, bold=bold, color=color)
    return p


def add_bullet(doc, text, *, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(4)
    if bold_prefix:
        r1 = p.add_run(bold_prefix)
        set_run(r1, size=11, bold=True, color=NAVY)
        r2 = p.add_run(text)
        set_run(r2, size=11, color=SLATE)
    else:
        r = p.add_run(text)
        set_run(r, size=11, color=SLATE)
    return p


def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        shade_cell(cell, HEADER_BG)
        cell_text(cell, h, bold=True, color=WHITE, size=10)
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            if r_idx % 2 == 1:
                shade_cell(cell, ROW_ALT)
            cell_text(cell, val, size=10)
    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Cm(w)
    doc.add_paragraph()
    return table


def add_callout(doc, title, body, bg="FFF7ED"):
    table = doc.add_table(rows=1, cols=1)
    cell = table.rows[0].cells[0]
    shade_cell(cell, bg)
    cell.text = ""
    p1 = cell.paragraphs[0]
    r = p1.add_run(title)
    set_run(r, size=11, bold=True, color=NAVY)
    p2 = cell.add_paragraph()
    r2 = p2.add_run(body)
    set_run(r2, size=10, color=SLATE)
    doc.add_paragraph()


def set_page(doc):
    section = doc.sections[0]
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)

    footer = section.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Uso interno — Kickoff do projeto  ·  Blindagem de Domínio  ·  30/08/2026  ·  Página ")
    set_run(run, size=8, color=MUTED)
    # page number field
    fld1 = OxmlElement("w:fldChar")
    fld1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld2 = OxmlElement("w:fldChar")
    fld2.set(qn("w:fldCharType"), "end")
    r2 = p.add_run()
    r2._r.append(fld1)
    r2._r.append(instr)
    r2._r.append(fld2)
    set_run(r2, size=8, color=MUTED)


def cover(doc):
    for _ in range(3):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("KICKOFF DO PROJETO")
    set_run(r, size=12, bold=True, color=GOLD)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Blindagem de Domínio")
    set_run(r, size=32, bold=True, color=NAVY)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Plataforma de autenticação de e-mail (SPF, DKIM e DMARC)")
    set_run(r, size=14, color=MUTED)

    add_p(doc, "")
    meta = [
        ("Tipo", "Documento operacional + especificação técnica"),
        ("Data", "30 de agosto de 2026"),
        ("Objetivo do dia", "Começar o projeto hoje — você destrava a operação; o desenvolvimento inicia a Fase 1"),
        ("Público", "Sócio / operação de e-mail e DNS / desenvolvimento / comercial"),
        ("Status", "Pronto para reunião e para execução — produto ainda não implementado"),
    ]
    add_table(doc, ["Campo", "Valor"], meta, col_widths=[4.5, 12.5])

    add_callout(
        doc,
        "Como usar este arquivo",
        "Parte A = o que VOCÊ faz hoje e o que eu preciso receber para programar sem bloqueio. "
        "Parte B = especificação para a equipe analisar (arquitetura, fases, riscos). "
        "Preencha o Anexo 1 (respostas) e me devolva — com isso eu começo o código da Fase 1 no mesmo dia.",
        bg="E8EEF4",
    )
    doc.add_page_break()


def parte_a(doc):
    add_heading_styled(doc, "Parte A — Começar hoje", 1)
    add_p(
        doc,
        "Você pediu para iniciar ainda hoje. Dá. O código da Fase 1 (scanner, painel, gerador de registros, "
        "login MSP/cliente) não depende de API do Registro.br. Depende de decisões e de um domínio de teste. "
        "A Fase 2 (CNAME hospedado + relatórios XML) depende de infra sua, que pode andar em paralelo.",
    )

    add_heading_styled(doc, "1. Divisão do trabalho neste primeiro dia", 2)
    add_table(
        doc,
        ["Quem", "Faz hoje", "Não precisa esperar"],
        [
            [
                "Você / operação",
                "Fechar as 10 decisões, escolher 2–3 domínios de teste, confirmar onde está o DNS de cada um, abrir caixa rua@ e zona dmarc.",
                "Não precisa publicar DMARC em cliente real hoje.",
            ],
            [
                "Você / comercial",
                "Definir nome do produto, URL do painel e se vende como add-on do e-mail.",
                "Não precisa de material de vendas pronto.",
            ],
            [
                "Eu (desenvolvimento)",
                "Começar Fase 1 com os defaults da seção 3, assim que o Anexo 1 voltar preenchido.",
                "Não preciso de senha de cliente nem de API do Registro.br para o MVP.",
            ],
        ],
        col_widths=[3.5, 7.5, 6.0],
    )

    add_heading_styled(doc, "2. O que você precisa fazer — checklist da sua parte", 2)
    add_p(doc, "Faça nesta ordem. Os itens 1 a 6 destravam o código hoje. Os itens 7 a 10 destravam a Fase 2 (medida de resultado).", bold=False)

    add_heading_styled(doc, "Bloco 1 — Decisões (30 a 45 minutos)", 3)
    add_p(doc, "Responda no Anexo 1 e me envie. Se não quiser discutir, aceite os defaults da seção 3 e escreva “pode seguir os defaults”.")

    add_table(
        doc,
        ["#", "Sua tarefa", "Por que importa", "Default se você não responder"],
        [
            ["D1", "Nome comercial + URL do painel (ex.: painel.suaempresa.com.br)", "Marca no login e nos e-mails do produto", "Blindagem de Domínio / painel local de desenvolvimento"],
            ["D2", "O produto nasce em repo separado ou neste monorepo?", "Evita misturar com o site de fianças", "Pasta isolada neste repo por enquanto"],
            ["D3", "MVP só Hostinger + Skymail + Registro.br?", "Limita o gerador e o checklist", "Sim; “outro DNS” vira checklist genérico"],
            ["D4", "Primeira entrega: só Fase 1 ou já esqueleto da Fase 2?", "Define o tamanho do primeiro sprint", "Fase 1 completa + telas/API stub da Fase 2"],
            ["D5", "Como o cliente entra? Você cria o usuário ou ele se cadastra?", "Auth e operação", "Você (admin) cria o login do cliente"],
            ["D6", "Delegação por CNAME ou TXT direto no cliente?", "Eficiência depois do dia 1", "CNAME, com fallback TXT"],
            ["D7", "Como receber XML do Gmail/Microsoft? IMAP ou HTTPS?", "Só Fase 2", "Caixa IMAP rua@ + worker"],
            ["D8", "Vende como add-on do e-mail ou SaaS avulso?", "Copy do painel e comercial", "Add-on mensal no pacote de e-mail"],
            ["D9", "Marca só da sua empresa ou white-label do cliente?", "Tema do portal", "Sua marca no MVP"],
            ["D10", "Quem autoriza p=reject no cliente? Só técnico?", "Segurança operacional", "Só operador interno; nunca automático no MVP"],
        ],
        col_widths=[1.5, 5.5, 5.0, 5.0],
    )

    add_heading_styled(doc, "Bloco 2 — Inventário da carteira (1 hora, pode ser incompleto)", 3)
    add_p(doc, "Abra uma planilha (ou use o Anexo 2) e liste pelo menos os 2–3 primeiros clientes que serão piloto. Ideal: 10–20 depois.")
    add_table(
        doc,
        ["Coluna", "O que preencher", "Como descobrir"],
        [
            ["Domínio", "cliente.com.br", "Contrato / painel"],
            ["E-mail", "Hostinger / Skymail / ambos / outro", "MX do domínio ou seu cadastro"],
            ["DNS autoritativo", "Registro.br / Hostinger / Skymail / Cloudflare / outro", "dig NS dominio.com.br  ou  whois"],
            ["Tem Registro.br com vocês?", "Sim / Não", "Conta do Registro.br"],
            ["Outros disparadores", "NF-e, RD Station, Mailchimp, ERP…", "Perguntar ao cliente ou ver depois no XML"],
            ["Pode ser piloto esta semana?", "Sim / Não", "Risco comercial"],
        ],
        col_widths=[4.5, 6.0, 6.5],
    )
    add_callout(
        doc,
        "Atalho para hoje",
        "Se a planilha inteira não sair hoje: me mande 2 domínios reais da carteira (pode ser o de vocês + 1 cliente que aceita teste) "
        "e diga o provedor de e-mail e onde está o DNS. Com isso eu calo o scanner e o gerador em casos reais.",
        bg="ECFDF5",
    )

    add_heading_styled(doc, "Bloco 3 — Domínio e e-mail de teste (hoje, 20 minutos)", 3)
    add_p(doc, "Eu preciso de um domínio que possamos “sujar” com registros de teste, sem risco de derrubar e-mail de cliente.")
    add_bullet(doc, " Preferência 1: um domínio de vocês (ex.: o da agência) com DNS que vocês controlam.")
    add_bullet(doc, " Preferência 2: um .com.br barato só para o produto (ex.: blindagemxxx.com.br).")
    add_p(doc, "Me envie:")
    add_bullet(doc, " o FQDN (nome do domínio);")
    add_bullet(doc, " se o DNS está no Registro.br, Hostinger ou Skymail;")
    add_bullet(doc, " se o e-mail desse domínio é Hostinger, Skymail ou nenhum.")
    add_p(doc, "Não me envie senha. Você publica os registros que eu gerar; eu só preciso consultar o DNS público.")

    add_heading_styled(doc, "Bloco 4 — Infra da Fase 2 (pode começar hoje, não trava a Fase 1)", 3)
    add_p(doc, "Isso é o que transforma o produto em “Sendmarc de verdade”. Pode ser aberto hoje e ficar pronto em 1–2 dias.")
    add_table(
        doc,
        ["Item", "O que criar", "Como", "Me devolver"],
        [
            [
                "Zona DNS nossa",
                "Um hostname para hospedar o DMARC dos clientes, ex.: dmarc.suaempresa.com.br",
                "No DNS que VOCÊS controlam, não no do cliente",
                "O hostname escolhido e confirmação de que vocês editam essa zona",
            ],
            [
                "Caixa de relatórios",
                "rua@reports.suaempresa.com.br (ou similar) que aceite muitos XMLs/dia",
                "Hostinger ou Skymail — caixa nova, sem uso pessoal",
                "O endereço. Senha só quando formos ligar o worker (Fase 2), via canal seguro",
            ],
            [
                "Subdomínio do painel",
                "painel.suaempresa.com.br (A/CNAME para o app)",
                "Pode esperar o primeiro deploy",
                "O hostname desejado",
            ],
        ],
        col_widths=[3.5, 5.0, 4.5, 4.0],
    )

    add_heading_styled(doc, "Bloco 5 — Acessos (não me envie senha de cliente)", 3)
    add_p(doc, "Para a Fase 1 eu NÃO preciso entrar no Registro.br, Hostinger ou Skymail de ninguém.")
    add_p(doc, "Vocês (operação) precisam ter, internamente:")
    add_bullet(doc, " Login no Registro.br dos clientes piloto (já têm).")
    add_bullet(doc, " hPanel Hostinger dos clientes com e-mail Hostinger.")
    add_bullet(doc, " Painel Skymail (Envios SMTP → gerar DKIM) dos clientes Skymail.")
    add_p(doc, "Quando a Fase 1 estiver no ar, o fluxo é: eu gero o TXT/CNAME no painel → vocês colam no DNS → eu revalido por lookup público.")

    add_heading_styled(doc, "Bloco 6 — Regras de segurança comercial (15 minutos)", 3)
    add_bullet(doc, " Combinar com a equipe: nenhum cliente real em p=reject sem checklist e autorização técnica.")
    add_bullet(doc, " Piloto desta semana só em p=none (monitorar), no máximo.")
    add_bullet(doc, " Copy comercial: nunca dizer “antivírus da caixa de entrada”.")
    add_bullet(doc, " Definir quem fala com o cliente piloto (CS ou técnico).")

    add_heading_styled(doc, "3. Defaults para eu começar HOJE sem reunião longa", 2)
    add_p(doc, "Se você responder só uma frase — “pode seguir os defaults” — eu assumo:")
    add_table(
        doc,
        ["Tema", "Default"],
        [
            ["Nome provisório", "Blindagem de Domínio"],
            ["Onde o código vive", "App isolado neste repositório (não mistura com o site de fianças)"],
            ["Provedores MVP", "Hostinger, Skymail, DNS Registro.br; resto = checklist genérico"],
            ["Entrega", "Fase 1 completa + esqueleto da Fase 2 (telas de relatório vazias, prontas para XML)"],
            ["Login", "Admin cria usuário do cliente"],
            ["Delegação", "CNAME, fallback TXT"],
            ["RUA", "IMAP depois; Fase 1 só deixa o campo rua= preparado"],
            ["Comercial", "Add-on do e-mail"],
            ["Marca", "A de vocês"],
            ["reject", "Somente operador interno, manual"],
        ],
        col_widths=[5.0, 12.0],
    )

    add_heading_styled(doc, "4. O que eu preciso receber — passo a passo", 2)
    add_p(doc, "Me devolva nesta ordem. Cada etapa desbloqueia um pedaço do desenvolvimento.")

    add_heading_styled(doc, "Passo 1 — Agora (e-mail / recado de 5 linhas)", 3)
    add_p(doc, "Texto mínimo que eu preciso:")
    add_p(
        doc,
        "1) Pode seguir os defaults?  (sim / segue o Anexo 1)\n"
        "2) Nome do produto e URL desejada do painel.\n"
        "3) Dois ou três domínios piloto: dominio — e-mail (Hostinger/Skymail) — DNS (Registro.br/Hostinger/Skymail).\n"
        "4) Um domínio de TESTE nosso (pode ser o da empresa).\n"
        "5) Hostname pensado para dmarc. e para rua@ (mesmo que a caixa ainda não exista).",
        size=11,
    )
    add_p(doc, "Com o Passo 1 eu começo: projeto, telas do dashboard, scanner DNS, gerador Hostinger/Skymail, login admin/cliente.", bold=True, color=NAVY)

    add_heading_styled(doc, "Passo 2 — Hoje à tarde (planilha curta)", 3)
    add_p(doc, "Anexo 2 com os pilotos. Se algum campo for “não sei”, deixe em branco — o scanner descobre NS e MX.")
    add_p(doc, "Com o Passo 2 eu calo o score nos domínios reais de vocês e ajusto seletores DKIM (hostingermail-a/b/c, Skymail).")

    add_heading_styled(doc, "Passo 3 — Hoje ou amanhã (infra Fase 2, sem senha)", 3)
    add_bullet(doc, " Criar a caixa rua@… e me dizer o endereço.")
    add_bullet(doc, " Confirmar que existe uma zona DNS que vocês editam para hospedar cli_XXXX.dmarc.suaempresa.com.br.")
    add_p(doc, "Com o Passo 3 eu deixo o CNAME e o rua= já gerados no assistente. O worker de XML entra quando a caixa existir.")

    add_heading_styled(doc, "Passo 4 — Quando a Fase 1 estiver no ar (operação, não desenvolvimento)", 3)
    add_p(doc, "Você (não eu) entra no DNS do domínio de TESTE e publica o que o painel mostrar. Depois clica em “Revalidar”.")
    add_p(doc, "Só então repetimos num cliente piloto, em p=none. Eu preciso que você me confirme: “publiquei, pode revalidar o domínio X”.")

    add_heading_styled(doc, "Passo 5 — Só na Fase 2 (canal seguro)", 3)
    add_p(doc, "Credencial SOMENTE da caixa rua@ (não de clientes), para o worker baixar os XML. Nunca cole senha no chat aberto se puder evitar — use o meio que vocês já usam para segredo.")
    add_p(doc, "Com o Passo 5 eu ligo ingest + gráficos de “quem tentou usar o domínio”.")

    add_heading_styled(doc, "5. O que eu NÃO preciso — não perca tempo com isso hoje", 2)
    add_bullet(doc, " Senha do Registro.br, Hostinger ou Skymail de cliente.")
    add_bullet(doc, " API EPP do Registro.br / credenciamento de registrador.")
    add_bullet(doc, " Certificado BIMI/VMC.")
    add_bullet(doc, " Texto jurídico final dos termos (dá para placeholder).")
    add_bullet(doc, " Logo final em todos os tamanhos (um PNG já serve).")
    add_bullet(doc, " Publicar p=reject em alguém.")

    add_heading_styled(doc, "6. Roteiro sugerido das próximas 8 horas", 2)
    add_table(
        doc,
        ["Quando", "Você", "Eu"],
        [
            ["Agora (15 min)", "Ler Parte A e responder o Passo 1 (5 linhas)", "Aguardo o Passo 1; se já mandou “seguir defaults”, inicio o app"],
            ["+1 h", "Listar 2–3 domínios piloto (Anexo 2)", "Scaffold do app, login, layout do dashboard de score"],
            ["+3 h", "Abrir caixa rua@ e anotar o endereço", "Scanner DNS + cards Imitação / Privacidade / Branding"],
            ["+5 h", "Confirmar zona dmarc. nossa", "Gerador Hostinger/Skymail + checklist Registro.br"],
            ["Ainda hoje (fim)", "Publicar registros só no domínio de TESTE quando eu mandar o TXT", "Revalidação + painel MSP da carteira piloto"],
            ["Amanhã", "Olhar o piloto interno e validar o comercial (add-on)", "Ajustes do que a equipe apontar no Word / no painel"],
        ],
        col_widths=[3.5, 7.0, 6.5],
    )

    add_heading_styled(doc, "7. Critério de “começamos hoje” (pronto / não pronto)", 2)
    add_table(
        doc,
        ["Pronto para eu programar a Fase 1", "Ainda não — me falta"],
        [
            ["“Pode seguir os defaults” OU Anexo 1 preenchido", "Nenhuma decisão e nenhum domínio de teste"],
            ["Pelo menos 1 domínio de teste nosso", "Só domínio de cliente em produção, sem autorização"],
            ["Provedor de e-mail dos pilotos (mesmo que aproximado)", "Nada — neste caso eu assumo Hostinger+Skymail no gerador"],
        ],
        col_widths=[8.5, 8.5],
    )
    add_callout(
        doc,
        "Pedido direto",
        "Me responda o Passo 1 neste mesmo fio. Com isso eu começo o código ainda hoje. "
        "O Word é para a equipe; o Passo 1 é para eu não ficar parado.",
        bg="FEF3C7",
    )
    doc.add_page_break()


def anexos_preenchiveis(doc):
    add_heading_styled(doc, "Anexo 1 — Respostas para devolver (copie e preencha)", 1)
    add_p(doc, "Pode responder “default” em qualquer linha.")

    add_table(
        doc,
        ["ID", "Pergunta", "Sua resposta"],
        [
            ["D1", "Nome do produto + URL do painel", ""],
            ["D2", "Repo isolado ou pasta neste repo?", ""],
            ["D3", "MVP só Hostinger + Skymail + Registro.br?", ""],
            ["D4", "Fase 1 só, ou Fase 1 + esqueleto Fase 2?", ""],
            ["D5", "Admin cria o usuário do cliente?", ""],
            ["D6", "CNAME ou TXT direto?", ""],
            ["D7", "RUA: IMAP (caixa) ou HTTPS?", ""],
            ["D8", "Add-on de e-mail ou SaaS avulso?", ""],
            ["D9", "Marca própria ou white-label?", ""],
            ["D10", "Quem autoriza p=reject?", ""],
            ["T1", "Domínio de TESTE (FQDN)", ""],
            ["T2", "DNS e e-mail desse domínio de teste", ""],
            ["T3", "Hostname dmarc. (Fase 2)", ""],
            ["T4", "Endereço rua@ (mesmo se a caixa ainda for criada)", ""],
            ["T5", "Pilotos: dominio | e-mail | DNS | pode testar?", "1) \n2) \n3) "],
            ["T6", "Logo (link ou “mando depois”)", ""],
        ],
        col_widths=[2.0, 7.5, 7.5],
    )

    add_heading_styled(doc, "Anexo 2 — Inventário de pilotos", 1)
    add_table(
        doc,
        ["Domínio", "E-mail", "DNS", "Registro.br?", "Outros envios", "Piloto?"],
        [["", "", "", "", "", ""], ["", "", "", "", "", ""], ["", "", "", "", "", ""], ["", "", "", "", "", ""]],
        col_widths=[3.5, 2.8, 2.8, 2.4, 3.2, 2.0],
    )
    doc.add_page_break()


def parte_b(doc):
    add_heading_styled(doc, "Parte B — Especificação técnica (para a equipe)", 1)
    add_p(
        doc,
        "Versão em Word do documento interno. Serve para análise. Não substitui a Parte A se o objetivo de hoje é começar.",
    )

    add_heading_styled(doc, "8. Resumo executivo", 2)
    add_p(
        doc,
        "Oferecer aos clientes de web e e-mail uma plataforma que (1) diagnostica o domínio, "
        "(2) implanta SPF/DKIM/DMARC, (3) mostra quem tentou usar o nome da empresa e "
        "(4) sobe a política até o bloqueio (p=reject) sem quebrar o e-mail legítimo.",
    )
    add_p(
        doc,
        "O bloqueio não é filtro de caixa de entrada. É o mesmo mecanismo da Sendmarc: Gmail, Outlook e Yahoo "
        "leem o DMARC e recusam a mensagem forjada. Nós não interceptamos o SMTP do criminoso.",
    )
    add_p(
        doc,
        "Posição: serviço gerenciado na carteira que já operamos (site + e-mail + Registro.br), "
        "não um clone público da Sendmarc. App próprio, separado do site de cartas de fiança.",
    )

    add_heading_styled(doc, "9. O problema", 2)
    add_p(
        doc,
        "O SMTP não verifica identidade. Qualquer pessoa pode enviar com From: financeiro@cliente.com.br "
        "(email spoofing / personificação de marca): boleto falso, phishing, queda de reputação e piora da entrega do e-mail real.",
    )
    add_table(
        doc,
        ["Protocolo", "Função", "Analogia"],
        [
            ["SPF", "Lista pública de IPs/servidores autorizados a enviar pelo domínio", "Autorização de remetente"],
            ["DKIM", "Assinatura criptográfica; chave pública no DNS", "Lacre de autenticidade"],
            ["DMARC", "O que fazer se SPF/DKIM falharem + relatórios", "Política de bloqueio + câmera"],
        ],
        col_widths=[3.0, 8.5, 5.5],
    )
    add_p(doc, "Sem DMARC, SPF e DKIM não bloqueiam spoofing. O botão de bloqueio é a política DMARC.", bold=True)

    add_heading_styled(doc, "10. O que o produto é e não é", 2)
    add_p(doc, "É", bold=True)
    add_bullet(doc, " Plataforma de cibersegurança de domínio corporativo, operada por nós.")
    add_bullet(doc, " Painel do cliente (score, política, tentativas) e painel MSP da carteira.")
    add_p(doc, "Não é", bold=True)
    add_bullet(doc, " Antivírus / filtro da caixa pessoal.")
    add_bullet(doc, " Proteção contra domínio parecido (lookalike) — outro produto.")
    add_bullet(doc, " Publicação automática no Registro.br (não há API pública de zona).")
    add_bullet(doc, " BIMI/VMC no MVP.")

    add_p(doc, "Pode prometer: implantar autenticação; mostrar quem usou o domínio (após rua=); chegar a reject com promoção segura; melhorar entrega do e-mail legítimo.")
    add_p(doc, "Não pode prometer: “antivírus da inbox”; bloqueio no mesmo dia com vários disparadores não mapeados; publicação 100% API no Registro.br.")

    add_heading_styled(doc, "11. Personas", 2)
    add_bullet(doc, " Operador interno: cadastra domínio, publica o que o painel gera, promove política.", bold_prefix="Operador. ")
    add_bullet(doc, " Cliente final: vê score e relatório em português; não edita DNS.", bold_prefix="Cliente. ")
    add_bullet(doc, " Comercial/CS: usa score e PDF mensal como upsell do e-mail.", bold_prefix="Comercial. ")

    add_heading_styled(doc, "12. O que copiamos da Sendmarc", 2)
    add_p(doc, "Eles não reescrevem o TXT _dmarc do cliente a cada mudança. Pedem CNAME:")
    add_p(doc, "_dmarc.cliente.com.br.  CNAME  cliente.dmarc.seudominio.com.br.", bold=True)
    add_p(doc, "A política vive no DNS deles. rua= aponta para a caixa deles. XML diário vira o painel.")
    add_p(doc, "Copiamos: CNAME + inbox RUA + parser + promoção segura. Não copiamos no MVP: BIMI, intel global, flattening maduro, white-glove enterprise.")

    add_heading_styled(doc, "13. Stack real (Hostinger, Skymail, Registro.br)", 2)
    add_p(doc, "Registro.br — acesso à conta ≠ API de zona. Só edita DNS se o NS for do próprio Registro.br (modo avançado). EPP é de registrador. Publicação inicial = passo humano. Não impede o produto se o CNAME for único.")
    add_p(doc, "Hostinger Email — registros típicos:", bold=True)
    add_p(
        doc,
        "MX: mx1/mx2.hostinger.com\n"
        "SPF: v=spf1 include:_spf.mail.hostinger.com ~all\n"
        "DKIM CNAME: hostingermail-a|b|c._domainkey → hostingermail-*.dkim.mail.hostinger.com\n"
        "Se o NS está na Hostinger, o hPanel já aplica MX+SPF+DKIM. DMARC costuma ser manual.",
        size=10,
    )
    add_p(doc, "Skymail — registros típicos:", bold=True)
    add_p(
        doc,
        "SPF: v=spf1 include:spf.skymail.net.br -all\n"
        "DKIM: gerado no painel (Envios SMTP); automático só se o DNS for deles.\n"
        "Há API admin de caixas; DKIM/DMARC não são o fluxo principal da API.",
        size=10,
    )

    add_table(
        doc,
        ["Tarefa", "Automático?", "Como"],
        [
            ["Score DNS de qualquer domínio", "Sim", "Lookups públicos"],
            ["Gerar SPF Hostinger e/ou Skymail", "Sim", "Templates; um único TXT"],
            ["DKIM Hostinger se DNS = Hostinger", "Quase", "Registros conhecidos / hPanel"],
            ["DKIM Skymail", "Semi", "Painel Skymail + gravar DNS"],
            ["Publicar _dmarc no Registro.br", "Manual (1×)", "Sem API de zona"],
            ["Mudar política depois do CNAME", "Sim", "Nosso DNS"],
            ["Ver tentativas de spoofing", "Sim após RUA", "Parser XML"],
        ],
        col_widths=[6.0, 3.5, 7.5],
    )
    add_p(doc, "Regra nº 1: o sistema resolve NS e diz ONDE publicar. Publicar no Registro.br com NS na Hostinger não muda nada.")

    add_heading_styled(doc, "14. Arquitetura", 2)
    add_p(
        doc,
        "Painel React (cliente + MSP) → API FastAPI/JWT → (a) worker de DNS, (b) policy engine, (c) ingest RUA. "
        "Banco MongoDB. Zona DNS nossa guarda o TXT DMARC. Mailbox rua@ recebe XML do Gmail/Microsoft/Yahoo.",
    )
    add_p(
        doc,
        "Stack alinhada ao repo atual: React + Tailwind + shadcn + Recharts; FastAPI; Motor/Mongo; JWT/bcrypt; dnspython; parser XML.",
    )
    add_p(doc, "Não misturar rotas do site de fianças com o produto de segurança.")

    add_heading_styled(doc, "15. Delegação DMARC", 2)
    add_p(doc, "1) ID estável por domínio, ex.: cli_8f3a")
    add_p(doc, "2) No DNS NOSSO:")
    add_p(doc, 'cli_8f3a.dmarc.empresa.com.br.  TXT  "v=DMARC1; p=none; rua=mailto:rua+cli_8f3a@reports.empresa.com.br; fo=1;"', size=10)
    add_p(doc, "3) No DNS do CLIENTE, uma vez:")
    add_p(doc, " _dmarc.cliente.com.br.  CNAME  cli_8f3a.dmarc.empresa.com.br.")
    add_p(doc, "4) Trocas de política = update só do passo 2. TTL 300–600s. CNAME e TXT _dmarc não coexistem. Fallback: TXT direto no cliente se o Registro.br recusar CNAME.")

    add_heading_styled(doc, "16. Relatórios DMARC (RUA)", 2)
    add_p(
        doc,
        "XML diário compactado (gz/zip), sem corpo de e-mail: IP, volume, From, SPF/DKIM, disposição. "
        "Proposta: mailbox + worker 5–15 min; endereço por tenant (rua+cli_8f3a@); idempotência org_name+report_id. "
        "Enriquecer IP com ASN e classificar hostinger | skymail | conhecido | desconhecido | suspeito. "
        "Sem RUA o produto só prova que a chave existe; com RUA prova que o golpe parou.",
    )

    add_heading_styled(doc, "17. Score proposto", 2)
    add_p(doc, "Imitação peso 70 (SPF único, SPF ≤10 lookups, DKIM, DMARC presente, enforcement). Privacidade peso 20 (TLS-RPT, MTA-STS). Branding peso 10 (BIMI — visível, pouco peso). NS/MX são selos, não score.")
    add_table(
        doc,
        ["Score", "Risco", "Texto ao cliente"],
        [
            ["80–100", "Baixo", "E-mails protegidos contra personificação."],
            ["50–79", "Moderado", "Há medidas; o bloqueio ainda não está completo."],
            ["0–49", "Alto", "O domínio pode ser usado em golpes."],
        ],
        col_widths=[3.0, 3.5, 10.5],
    )
    add_p(doc, "Calibrar com 10–20 domínios reais antes de tratar o número como verdade absoluta.")

    add_heading_styled(doc, "18. SPF — regras que evitam incidente", 2)
    add_bullet(doc, " Um único v=spf1 no apex. Dois TXT = falha.")
    add_bullet(doc, " Máximo 10 lookups. Hostinger + Skymail + NF-e + marketing estoura. Flattening = Fase 3.")
    add_bullet(doc, " ~all no começo; -all só com relatórios limpos; +all proibido no gerador.")
    add_bullet(doc, " Ambos provedores: v=spf1 include:_spf.mail.hostinger.com include:spf.skymail.net.br ~all")
    add_bullet(doc, " Nunca apagar include “desconhecido” sem olhar o XML.")

    add_heading_styled(doc, "19. Promoção segura de política", 2)
    add_p(doc, "none → quarantine → reject, com rollback.")
    add_p(doc, "Quarantine: rua= nosso com ≥7 dias de XML; ≥95% do volume conhecido passando SPF ou DKIM; nenhum conhecido falhando em série.")
    add_p(doc, "Reject: janela ≥14 dias em quarantine (ou 21 em none se volume baixo); zero conhecido falhando em 7 dias; confirmação humana do operador.")
    add_p(doc, "Reject no escuro é o maior risco do produto.", bold=True)

    add_heading_styled(doc, "20. Dados e API (resumo)", 2)
    add_p(doc, "Coleções: tenants, users (admin|operator|client), domains, scans, dmarc_reports, dmarc_rows.")
    add_p(doc, "API /api/shield: login, msp/domains, scan, setup, setup/verify, senders, policy. Scanner público com rate-limit se o comercial quiser um “cole o domínio”.")

    add_heading_styled(doc, "21. Telas", 2)
    add_p(doc, "Cliente: score circular, risco, selos NS/MX, cards Imitação / Privacidade / Branding, relatório detalhado, gráfico 30 dias (Fase 2).")
    add_p(doc, "MSP: carteira, filas (publicar CNAME, validar DKIM Skymail, candidato a quarantine), detalhe + política + auditoria.")
    add_p(doc, "Copy: p=none = Apenas monitoramento; p=quarantine = Quarentena (spam); p=reject = Bloqueio total.")

    add_heading_styled(doc, "22. Fases", 2)
    add_table(
        doc,
        ["Fase", "Entrega", "Já vende?"],
        [
            ["1", "Auth, cadastro, scanner ao vivo, score, gerador, checklist DNS, revalidação, MSP", "Sim — implantação assistida"],
            ["2", "CNAME nosso, mailbox RUA, parser, gráficos, policy engine, alertas", "Sim — “Sendmarc” de verdade"],
            ["3", "Flattening SPF, MTA-STS hospedado, white-label, PDF mensal, Cloudflare no checklist", "Escala"],
        ],
        col_widths=[2.0, 11.5, 3.5],
    )

    add_heading_styled(doc, "23. Infra da empresa (não é código)", 2)
    add_table(
        doc,
        ["Item", "Para quê", "Sem isso"],
        [
            ["Zona DNS nossa", "Hospedar política DMARC", "Cada reject = edição no cliente"],
            ["Caixa rua@", "Relatórios", "Painel sem prova de bloqueio"],
            ["Acessos que já existem", "Publicar o dia 1", "Checklist — aceitável"],
            ["Marca e URL", "Portal do cliente", "Produto genérico"],
        ],
        col_widths=[4.5, 6.5, 6.0],
    )

    add_heading_styled(doc, "24. Segurança, LGPD, riscos", 2)
    add_bullet(doc, " RUA não traz corpo de e-mail; ainda assim é metadado. Isolar por tenant. Retenção sugerida: 13 meses.")
    add_bullet(doc, " MVP: não guardar senha de Registro.br/Hostinger/Skymail no sistema.")
    add_bullet(doc, " Rate-limit no scanner. Log de quem mudou política. Worker RUA só aceita XML/gz/zip.")
    add_table(
        doc,
        ["Risco", "Mitigação"],
        [
            ["reject cedo demais", "Checklist + humano + rollback"],
            ["Dois TXT SPF", "Scanner laranja; gerador substitui"],
            [">10 lookups", "Contador; flattening na Fase 3"],
            ["Publicar no DNS errado", "Detectar NS e travar o checklist"],
            ["CNAME recusado no Registro.br", "Fallback TXT"],
            ["Cliente acha que é antivírus", "Copy fixo no comercial e no onboarding"],
        ],
        col_widths=[6.5, 10.5],
    )

    add_heading_styled(doc, "25. Testes de aceite (quando houver código)", 2)
    add_bullet(doc, " Domínio com SPF+DKIM+DMARC → Imitação alta; só SPF/DKIM sem DMARC → risco alto.")
    add_bullet(doc, " Dois SPF → laranja. Gerador Hostinger = 3 CNAMEs oficiais. Gerador “ambos” = um SPF com dois includes.")
    add_bullet(doc, " NS Registro.br vs Hostinger escolhe o checklist certo. Tenant isolado.")
    add_bullet(doc, " Fase 2: dig TXT _dmarc segue o CNAME; política reflete após TTL; XML fixture idempotente; reject bloqueado se checklist falhar.")

    add_heading_styled(doc, "26. Perguntas extras para a reunião (além do Anexo 1)", 2)
    add_bullet(doc, " Quantos domínios hoje estão em DNS Registro.br vs Hostinger vs Skymail?")
    add_bullet(doc, " Quantos usam os dois e-mails ou NF-e/marketing além do corporativo?")
    add_bullet(doc, " Já existe domínio/caixa para dmarc. e rua@ ou é compra nova?")
    add_bullet(doc, " O comercial vende “segurança” ou “e-mail que não cai em spam”?")

    add_heading_styled(doc, "27. Conclusão", 2)
    add_p(
        doc,
        "É possível e vale a pena como serviço gerenciado da carteira, com três condições: "
        "(1) não tentar ser Sendmarc global; (2) aceitar o dia 1 humano no Registro.br e ganhar eficiência no CNAME + relatórios; "
        "(3) nunca ligar p=reject sem XML e sem mapear remetentes.",
    )
    add_p(
        doc,
        "O protocolo é barato de consultar. O que o cliente paga todo mês é operar a política sem derrubar o e-mail e provar que o spoofing parou.",
    )
    add_p(
        doc,
        "Próximo passo: devolver o Anexo 1 (ou “seguir defaults”) + 1 domínio de teste. Com isso o desenvolvimento da Fase 1 começa no mesmo dia.",
        bold=True,
        color=NAVY,
    )

    add_heading_styled(doc, "28. Referências", 2)
    add_p(
        doc,
        "RFC 7208 SPF · RFC 6376 DKIM · RFC 7489 DMARC · RFC 8461 MTA-STS · RFC 8460 TLS-RPT · "
        "Hostinger include:_spf.mail.hostinger.com e hostingermail-a/b/c._domainkey · "
        "Skymail include:spf.skymail.net.br · Registro.br zona avançada sem API pública · "
        "Sendmarc: CNAME _dmarc + RUA hospedado (modelo, não marca).",
        size=10,
    )
    add_p(doc, "Nenhuma implementação deste produto existia no repositório no momento desta redação.", size=10, color=MUTED)


def main():
    out = Path("/workspace/docs/Blindagem-de-Dominio-Kickoff-e-Especificacao.docx")
    out.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()
    set_page(doc)
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    style.font.color.rgb = SLATE

    cover(doc)
    parte_a(doc)
    anexos_preenchiveis(doc)
    parte_b(doc)

    doc.save(out)
    print(f"Wrote {out} ({out.stat().st_size} bytes)")


if __name__ == "__main__":
    main()
